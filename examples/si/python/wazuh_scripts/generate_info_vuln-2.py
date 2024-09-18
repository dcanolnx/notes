#!/usr/bin/env python3
import logging, os, configparser, requests, paramiko, csv, sys
from datetime import datetime, timedelta
from collections import Counter, defaultdict
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
from collections import defaultdict
from docx import Document
from docx.shared import Pt
import matplotlib.pyplot as plt
from docx.shared import Inches
import io
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt
from docx.shared import Inches, Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx import Document

from collections import Counter
from docx.shared import Inches, Pt, Cm
from docx.enum.section import WD_ORIENT
from docx import Document
import matplotlib.pyplot as plt
import io



config_file_path = "./wazuh.config"


def set_cvss3_background(cell, score):
    if score >= 8:
        set_cell_background(cell, "FF0000")  # Rojo
    elif score >= 7:
        set_cell_background(cell, "FFA500")  # Naranja
    else:
        set_cell_background(cell, "FFFFFF")  # Blanco (o cualquier otro color que prefieras para los puntajes más bajos)

def load_config(config_file_path):
    config = configparser.ConfigParser()
    config.read(config_file_path)
    return config

def get_auth_token(config):
    api_url = "https://wazuh-aws.int.stratio.com:55000/security/user/authenticate?raw=true"
    username = config['WAZUH']["USER"]
    password = config['WAZUH']["PASSWORD"]
    response = requests.get(api_url, auth=(username, password), verify=False)
    
    if response.status_code == 200:
        return response.text
    else:
        print(f"Request wazuh({api_url}) fail HTTP response is: {response.status_code}")
        return None

def get_agent_ids_from_group(config, group_name, token):
    url = f"https://wazuh-aws.int.stratio.com:55000/groups/{group_name}/agents"
    response = requests.get(url, verify=False, headers={'Authorization': 'Bearer ' + token})
    
    if response.status_code == 200:
        data = response.json()
        # La clave para el hostname según la salida de la API es 'name'
        agents_info = [(item['id'], item['os']['name'], item['name']) for item in data['data']['affected_items']]
        return agents_info
    else:
        print(f"Failed to get agents from group {group_name}. HTTP response is: {response.status_code}")
        return []
        
def get_vulnerabilities_by_agent(config, agent_id, token):
    url = f"https://wazuh-aws.int.stratio.com:55000/vulnerability/{agent_id}"
    response = requests.get(url, verify=False, headers={'Authorization': 'Bearer ' + token})
    
    if response.status_code == 200:
        data = response.json()
        vulnerabilities = data['data'].get('affected_items') or data['data'].get('items', [])
        return vulnerabilities
    else:
        print(f"Failed to get vulnerabilities for agent {agent_id}. HTTP response is: {response.status_code}")
        return []


def adjust_table_columns(table, widths):
    for i, column in enumerate(table.columns):
        set_column_width(column, widths[i])

# Función para configurar el fondo de una celda
def set_cell_background(cell, color_str):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_str))
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Función para establecer el ancho de las celdas de una fila
def set_row_cell_widths(row, widths):
    for cell, width in zip(row.cells, widths):
        cell.width = width
def set_severity_background(cell, severity):
    if severity == 'High':
        set_cell_background(cell, "FF0000")  # Rojo
    elif severity == 'Medium':
        set_cell_background(cell, "FFA500")  # Naranja
    elif severity in ['Untriaged', 'Low']:
        set_cell_background(cell, "FFFF00")  # Amarillo

def set_margins(section, top, bottom, left, right):
    section.top_margin = top
    section.bottom_margin = bottom
    section.left_margin = left
    section.right_margin = right

def compile_vulnerabilities_by_distro(vulnerabilities_by_agent, agents_info):
    vulnerabilities_by_distro = defaultdict(lambda: defaultdict(list))
    # Convertir agents_info a un diccionario para acceso directo
    agents_dict = {agent[0]: {"distro": agent[1], "hostname": agent[2]} for agent in agents_info}
    for agent_id, vulnerabilities in vulnerabilities_by_agent.items():
        # Utilizar el diccionario para obtener la distro y hostname
        agent_data = agents_dict.get(agent_id, {'distro': 'Unknown', 'hostname': 'Unknown Hostname'})
        distro = agent_data['distro']
        for vulnerability in vulnerabilities:
            vulnerability_data = {
                'name': vulnerability.get('name'),
                'severity': vulnerability.get('severity'),
                'cvss3_score': vulnerability.get('cvss3_score'),
                'cve': vulnerability.get('cve'),
                # Incluir más campos según estén disponibles y sean relevantes
            }
            vulnerabilities_by_distro[distro][vulnerability['name']].append(vulnerability_data)
    return vulnerabilities_by_distro

def generate_agent_reports(vulnerabilities_by_agent, agent_id):
    vulnerabilities = vulnerabilities_by_agent.get(agent_id, [])
    hostname = "hostname_unknown"  # Esto es un valor predeterminado en caso de que no tengamos el hostname

    # Aquí crearías el documento como antes
    doc = Document()
    section = doc.sections[-1]
    set_margins(section, Cm(0), Cm(0), Cm(1), Cm(1))
    section.orientation = WD_ORIENT.LANDSCAPE

    # Título del informe
    doc.add_heading(f'Vulnerabilidades para el agente: {agent_id} - {hostname}', level=1)

    # Tabla basada en Componentes
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    headers = ['Componente', 'Conteo', 'Severidad', 'CVSS3 Score', 'CVE']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        set_cell_background(hdr_cells[i], "4169E1")

    # Llenar la tabla con los datos de vulnerabilidad
    sorted_vulnerabilities = sorted(vulnerabilities, key=lambda x: x['severity'], reverse=True)
    for vuln in sorted_vulnerabilities:
        row = table.add_row()
        row.cells[0].text = vuln['name']
        row.cells[1].text = '1'
        set_severity_background(row.cells[2], vuln['severity'])
        row.cells[2].text = vuln['severity']
        row.cells[3].text = str(vuln['cvss3_score'])
        set_cvss3_background(row.cells[3], vuln['cvss3_score'])
        row.cells[4].text = vuln.get('cve', 'N/A')


    # Crear un gráfico de las 10 componentes afectados más comunes
    vulnerability_counts = Counter(vulnerability['name'] for vulns in vulnerabilities.values() for vulnerability in vulns)
    common_vulns = vulnerability_counts.most_common(10)
    vuln_names = [vuln[0] for vuln in common_vulns]
    counts = [vuln[1] for vuln in common_vulns]        
    plt.figure(figsize=(10, 6))
    plt.barh(vuln_names, counts, color='orange')
    plt.xlabel('Número de ocurrencias')
    plt.title('Top 10 componentes afectados más comunes')
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    doc.add_paragraph()
    doc.add_picture(buf, width=section.page_width - section.left_margin - section.right_margin)
    buf.close()
    plt.close()



    # Crear un gráfico de pastel de severidades
    severity_counts = Counter(vuln['severity'] for vulns in vulnerabilities.values() for vuln in vulns)
    labels = severity_counts.keys()
    sizes = severity_counts.values()
    plt.figure(figsize=(6, 6))
    plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
    plt.axis('equal')
    plt.title('Distribución de Vulnerabilidades por Severidad')
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    doc.add_paragraph()
    doc.add_picture(buf, width=section.page_width - section.left_margin - section.right_margin)
    buf.close()
    plt.close()


    # Añadir una nueva tabla basada en CVEs
    doc.add_page_break()
    doc.add_heading('Vulnerabilidades por CVE', level=1)
    cve_table = doc.add_table(rows=1, cols=5)
    cve_hdr_cells = cve_table.rows[0].cells
    headers = ['CVE', 'Conteo', 'Severidad', 'CVSS3 Score', 'Componente']
    for i, header in enumerate(headers):
        cve_hdr_cells[i].text = header
        cve_hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        cve_hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        cve_hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        set_cell_background(cve_hdr_cells[i], "4169E1")

    # Agrupar vulnerabilidades por CVE y ordenar por conteo
    cve_vulnerabilities = defaultdict(list)
    for vuln_list in vulnerabilities.values():
        for vuln in vuln_list:
            cve_vulnerabilities[vuln.get('cve', 'N/A')].append(vuln)
    sorted_cve_vulnerabilities = sorted(cve_vulnerabilities.items(), key=lambda x: len(x[1]), reverse=True)
    for cve, vuln_list in sorted_cve_vulnerabilities:
        count = len(vuln_list)
        example_vuln = vuln_list[0]  # Ejemplo de vulnerabilidad para obtener detalles
        row = cve_table.add_row()
        row.cells[0].text = cve
        row.cells[1].text = str(count)
        set_severity_background(row.cells[2], example_vuln['severity'])
        row.cells[2].text = example_vuln['severity']
        row.cells[3].text = str(example_vuln['cvss3_score'])
        set_cvss3_background(row.cells[3], example_vuln['cvss3_score'])
        row.cells[4].text = example_vuln['name']  # Asignar el nombre del componente a la última celda



    # (El código de llenado de la tabla CVE se mantiene igual)
    # Añadir gráficos basados en CVEs
    # Crear un gráfico de barras de los 10 CVEs más comunes
    cve_counts = Counter(vuln['cve'] for vulns in vulnerabilities.values() for vuln in vulns if 'cve' in vuln)
    common_cves = cve_counts.most_common(10)
    cve_names = [cve[0] for cve in common_cves]
    counts = [cve[1] for cve in common_cves]
    plt.figure(figsize=(10, 6))
    plt.barh(cve_names, counts, color='salmon')
    plt.xlabel('Número de ocurrencias')
    plt.title('Top 10 CVEs más comunes')
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    # Añadir el gráfico al documento
    doc.add_paragraph()
    doc.add_picture(buf, width=section.page_width - section.left_margin - section.right_margin)
    buf.close()
    plt.close()
    # Gráfico de pastel de severidades para CVEs
    plt.figure(figsize=(6, 6))
    plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
    plt.axis('equal')
    plt.title('Distribución de Severidades por CVE')
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    doc.add_paragraph()
    doc.add_picture(buf, width=section.page_width - section.left_margin - section.right_margin)
    buf.close()
    plt.close()
    # Añadir la lista de agentes al final del documento en una tabla
    doc.add_page_break()
    doc.add_heading('Lista de Agentes', level=1)
    agent_table = doc.add_table(rows=1, cols=2)
    agent_table.style = 'Table Grid'
    agent_hdr_cells = agent_table.rows[0].cells
    agent_hdr_cells[0].text = 'ID'
    agent_hdr_cells[1].text = 'Hostname'
    for agent in filtered_agents_info:
        agent_row_cells = agent_table.add_row().cells
        agent_row_cells[0].text = agent[0]
        agent_row_cells[1].text = agent[1]

    
    # Guardar el documento con un nombre de archivo único para cada agente
    doc_file = f"vulnerabilities_agent_{agent_id}_{hostname}.docx"
    doc.save(doc_file)
    print(f"Archivo {doc_file} creado con éxito para el agente {agent_id} - {agent_data['hostname']}.")


def generate_distro_reports(vulnerabilities_by_distro, agents_info):
    for distro, vulnerabilities in vulnerabilities_by_distro.items():
        # Filtrar agents_info para incluir solo los agentes de la distribución actual
        filtered_agents_info = [(agent[0], agent[2]) for agent in agents_info if agent[1] == distro]
        agent_details = ', '.join([f"ID: {agent[0]}, Hostname: {agent[1]}" for agent in filtered_agents_info])

        # Crear un nuevo documento Word para cada distribución
        doc = Document()
        section = doc.sections[-1]
        set_margins(section, Cm(1), Cm(1), Cm(1), Cm(1))
        section.orientation = WD_ORIENT.LANDSCAPE
        
        # Agregar detalles de los agentes y el título
        doc.add_heading(f'Vulnerabilidades para la distribución: {distro}', level=1)
        
        # Añadir una tabla al documento con los encabezados actualizados
        # Tabla basada en Componentes
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        headers = ['Componente', 'Conteo', 'Severidad', 'CVSS3 Score', 'CVE']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            set_cell_background(hdr_cells[i], "4169E1")  # Un azul más fuerte

        # Ordenar por Conteo y llenar la tabla con los datos de vulnerabilidad
        sorted_vulnerabilities = sorted(vulnerabilities.items(), key=lambda x: len(x[1]), reverse=True)
        for vulnerability_name, vuln_list in sorted_vulnerabilities:
            count = len(vuln_list)
            vuln_details = vuln_list[0]  # Ejemplo de vulnerabilidad para obtener detalles
            row = table.add_row()
            row.cells[0].text = vulnerability_name
            row.cells[1].text = str(count)
            set_severity_background(row.cells[2], vuln_details['severity'])
            row.cells[2].text = vuln_details['severity']
            row.cells[3].text = str(vuln_details['cvss3_score'])
            set_cvss3_background(row.cells[3], vuln_details['cvss3_score'])
            row.cells[4].text = vuln_details.get('cve', 'N/A')


        # Crear un gráfico de las 10 componentes afectados más comunes
        vulnerability_counts = Counter(vulnerability['name'] for vulns in vulnerabilities.values() for vulnerability in vulns)
        common_vulns = vulnerability_counts.most_common(10)
        vuln_names = [vuln[0] for vuln in common_vulns]
        counts = [vuln[1] for vuln in common_vulns]
        
        plt.figure(figsize=(10, 6))
        plt.barh(vuln_names, counts, color='orange')
        plt.xlabel('Número de ocurrencias')
        plt.title('Top 10 componentes afectados más comunes')
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        doc.add_paragraph()
        doc.add_picture(buf, width=section.page_width - section.left_margin - section.right_margin)
        buf.close()
        plt.close()



        # Crear un gráfico de pastel de severidades
        severity_counts = Counter(vuln['severity'] for vulns in vulnerabilities.values() for vuln in vulns)
        labels = severity_counts.keys()
        sizes = severity_counts.values()

        plt.figure(figsize=(6, 6))
        plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
        plt.axis('equal')
        plt.title('Distribución de Vulnerabilidades por Severidad')

        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        doc.add_paragraph()
        doc.add_picture(buf, width=section.page_width - section.left_margin - section.right_margin)
        buf.close()
        plt.close()


        # Añadir una nueva tabla basada en CVEs
        doc.add_page_break()
        doc.add_heading('Vulnerabilidades por CVE', level=1)
        cve_table = doc.add_table(rows=1, cols=5)
        cve_hdr_cells = cve_table.rows[0].cells
        headers = ['CVE', 'Conteo', 'Severidad', 'CVSS3 Score', 'Componente']
        for i, header in enumerate(headers):
            cve_hdr_cells[i].text = header
            cve_hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            cve_hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            cve_hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            set_cell_background(cve_hdr_cells[i], "4169E1")

        # Agrupar vulnerabilidades por CVE y ordenar por conteo
        cve_vulnerabilities = defaultdict(list)
        for vuln_list in vulnerabilities.values():
            for vuln in vuln_list:
                cve_vulnerabilities[vuln.get('cve', 'N/A')].append(vuln)

        sorted_cve_vulnerabilities = sorted(cve_vulnerabilities.items(), key=lambda x: len(x[1]), reverse=True)
        for cve, vuln_list in sorted_cve_vulnerabilities:
            count = len(vuln_list)
            example_vuln = vuln_list[0]  # Ejemplo de vulnerabilidad para obtener detalles
            row = cve_table.add_row()
            row.cells[0].text = cve
            row.cells[1].text = str(count)
            set_severity_background(row.cells[2], example_vuln['severity'])
            row.cells[2].text = example_vuln['severity']
            row.cells[3].text = str(example_vuln['cvss3_score'])
            set_cvss3_background(row.cells[3], example_vuln['cvss3_score'])
            row.cells[4].text = example_vuln['name']  # Asignar el nombre del componente a la última celda



        # (El código de llenado de la tabla CVE se mantiene igual)

        # Añadir gráficos basados en CVEs
        # Crear un gráfico de barras de los 10 CVEs más comunes
        cve_counts = Counter(vuln['cve'] for vulns in vulnerabilities.values() for vuln in vulns if 'cve' in vuln)
        common_cves = cve_counts.most_common(10)
        cve_names = [cve[0] for cve in common_cves]
        counts = [cve[1] for cve in common_cves]

        plt.figure(figsize=(10, 6))
        plt.barh(cve_names, counts, color='salmon')
        plt.xlabel('Número de ocurrencias')
        plt.title('Top 10 CVEs más comunes')
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)

        # Añadir el gráfico al documento
        doc.add_paragraph()
        doc.add_picture(buf, width=section.page_width - section.left_margin - section.right_margin)
        buf.close()
        plt.close()


        # Gráfico de pastel de severidades para CVEs
        plt.figure(figsize=(6, 6))
        plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
        plt.axis('equal')
        plt.title('Distribución de Severidades por CVE')
        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        doc.add_paragraph()
        doc.add_picture(buf, width=section.page_width - section.left_margin - section.right_margin)
        buf.close()
        plt.close()
        # Añadir la lista de agentes al final del documento en una tabla
        doc.add_page_break()
        doc.add_heading('Lista de Agentes', level=1)
        agent_table = doc.add_table(rows=1, cols=2)
        agent_table.style = 'Table Grid'
        agent_hdr_cells = agent_table.rows[0].cells
        agent_hdr_cells[0].text = 'ID'
        agent_hdr_cells[1].text = 'Hostname'
        for agent in filtered_agents_info:
            agent_row_cells = agent_table.add_row().cells
            agent_row_cells[0].text = agent[0]
            agent_row_cells[1].text = agent[1]

        # Guardar el documento
        doc_file = f"vulnerabilities_distro_{distro}.docx"
        doc.save(doc_file)
        print(f"Archivo {doc_file} creado con éxito para la distribución {distro}.")



def generate_csv_report(vulnerabilities, filename):
    with open(filename, mode='w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(['Agente', 'Vulnerabilidad', 'Severidad', 'CVSS2 Score', 'CVSS3 Score', 'CVE'])
        for agent_id, vulns in vulnerabilities.items():
            for vuln in vulns:
                writer.writerow([
                    agent_id,
                    vuln['name'],
                    vuln['severity'],
                    vuln['cvss2_score'],
                    vuln['cvss3_score'],
                    vuln['cve']
                ])

def main():
    print("El Script tardará unos minutos en generar todos los informes...")
    if os.path.isfile(config_file_path):
        config = load_config(config_file_path)
        token = get_auth_token(config)
        if token:
            group_name = sys.argv[1] if len(sys.argv) > 1 else 'Servers_Linux'
            agents_info = get_agent_ids_from_group(config, group_name, token)

            vulnerabilities_by_agent = {}
            for agent_id, distro, hostname in agents_info:
                vulnerabilities = get_vulnerabilities_by_agent(config, agent_id, token)
                vulnerabilities_by_agent[agent_id] = vulnerabilities
            """
            # Aquí llamamos a la función para generar los informes de los agentes
            for agent_id, distro, hostname in agents_info:
                generate_agent_reports(vulnerabilities_by_agent, agent_id)
            """
            vulnerabilities_by_distro = compile_vulnerabilities_by_distro(vulnerabilities_by_agent, agents_info)
            

            # Crear documentos para cada tipo de informe
            doc_distro = Document()
            doc_agent = Document()
            doc_group = Document()
            
            # Generar los informes
            generate_distro_reports(vulnerabilities_by_distro, agents_info)
            

            #generate_group_report(vulnerabilities_by_agent, agents_info)

            # Guardar los documentos
            doc_distro.save(f"distro_vulnerabilities_{group_name}.docx")
            doc_agent.save(f"agent_vulnerabilities_{group_name}.docx")
            doc_group.save(f"group_vulnerabilities_{group_name}.docx")

            #print(f"Archivo {doc_file} creado con éxito para el grupo {group_name}.")

            csv_file = f"vulnerabilities_report_{group_name}.csv"
            generate_csv_report(vulnerabilities_by_agent, csv_file)
            print(f"Archivo {csv_file} creado con éxito para el grupo {group_name}.")

if __name__ == "__main__":
    main()
